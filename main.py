from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from google.cloud import storage
import psycopg2
from psycopg2.extras import RealDictCursor
import requests
import os
import tempfile
from typing import Optional
from urllib.parse import unquote
import re

app = FastAPI(title="Audio Migration API", version="1.0.0")

# Database configuration
DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://postgres:%40%21ceXpert%2C9%2F11@db.ycnqrnacilcwqqgelaod.supabase.co:5432/postgres?sslmode=require"
)

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://ycnqrnacilcwqqgelaod.supabase.co")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")  # You'll need to provide this
BUCKET_NAME = os.getenv("BUCKET_NAME", "audio-files")  # Default bucket name

# GCS Bucket name
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "clarodele-mvp-content")


def get_db_connection():
    """Create a database connection"""
    return psycopg2.connect(DATABASE_URL)


def extract_google_drive_id(url: str) -> Optional[str]:
    """Extract file ID from Google Drive URL"""
    patterns = [
        r'/file/d/([a-zA-Z0-9_-]+)',
        r'id=([a-zA-Z0-9_-]+)',
        r'/open\?id=([a-zA-Z0-9_-]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def download_from_google_drive(file_id: str, destination: str) -> bool:
    """Download file from Google Drive using direct download link"""
    try:
        # Google Drive direct download URL
        url = f"https://drive.google.com/uc?export=download&id={file_id}"
        
        session = requests.Session()
        response = session.get(url, stream=True)
        
        # Handle large files with confirmation token
        for key, value in response.cookies.items():
            if key.startswith('download_warning'):
                params = {'export': 'download', 'id': file_id, 'confirm': value}
                response = session.get(url, params=params, stream=True)
                break
        
        # Save file
        with open(destination, 'wb') as f:
            for chunk in response.iter_content(32768):
                if chunk:
                    f.write(chunk)
        
        return True
    except Exception as e:
        print(f"Error downloading from Google Drive: {str(e)}")
        return False


def upload_to_gcs(local_file_path: str, destination_blob_name: str) -> Optional[str]:
    """Upload file to Google Cloud Storage and return public URL"""
    try:
        # Initialize client - will use Application Default Credentials
        # For local: run `gcloud auth application-default login`
        # For production: uses GOOGLE_APPLICATION_CREDENTIALS
        client = storage.Client()
        bucket = client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(destination_blob_name)
        
        # Upload file
        blob.upload_from_filename(local_file_path)
        
        # Make the blob publicly accessible (optional)
        # blob.make_public()
        
        # Return the public URL
        return f"gs://{GCS_BUCKET_NAME}/{destination_blob_name}"
        # Or return public URL: return blob.public_url
    except Exception as e:
        print(f"Error uploading to GCS: {str(e)}")
        return None


def update_bucket_url(conn, row_id: int, bucket_url: str, table_name: str = "listening_tarea1_set", id_column: str = "tarea1_set_id"):
    """Update the bucket_url column for a specific row"""
    cursor = conn.cursor()
    try:
        query = f"UPDATE {table_name} SET bucket_url = %s WHERE {id_column} = %s"
        cursor.execute(query, (bucket_url, row_id))
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        cursor.close()


@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Root endpoint with HTML interface"""
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Audio Migration - Cloud Run</title>
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif;
                max-width: 800px;
                margin: 50px auto;
                padding: 20px;
                background: #f5f5f5;
            }
            .container {
                background: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            h1 {
                color: #333;
                border-bottom: 3px solid #4CAF50;
                padding-bottom: 10px;
            }
            .status-box {
                background: #e3f2fd;
                border-left: 4px solid #2196F3;
                padding: 15px;
                margin: 20px 0;
                border-radius: 4px;
            }
            .button {
                display: inline-block;
                padding: 12px 24px;
                margin: 10px 5px;
                background: #4CAF50;
                color: white;
                text-decoration: none;
                border-radius: 5px;
                font-weight: bold;
                transition: background 0.3s;
            }
            .button:hover {
                background: #45a049;
            }
            .button.secondary {
                background: #2196F3;
            }
            .button.secondary:hover {
                background: #0b7dda;
            }
            .endpoint {
                background: #f9f9f9;
                padding: 10px;
                margin: 5px 0;
                border-radius: 4px;
                font-family: monospace;
            }
            .loading {
                display: none;
                text-align: center;
                padding: 20px;
            }
            .result {
                margin-top: 20px;
                padding: 15px;
                background: #f9f9f9;
                border-radius: 5px;
                white-space: pre-wrap;
                font-family: monospace;
                font-size: 12px;
                max-height: 400px;
                overflow-y: auto;
            }
            .success { border-left: 4px solid #4CAF50; }
            .error { border-left: 4px solid #f44336; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🎵 Audio Migration System</h1>
            
            <div class="status-box">
                <strong>📊 Status:</strong> System Ready<br>
                <strong>🗄️ Database:</strong> Supabase PostgreSQL<br>
                <strong>☁️ Storage:</strong> Google Cloud Storage<br>
                <strong>📦 Bucket:</strong> clarodele-mvp-content
            </div>

            <h2>Quick Actions</h2>
            
            <a href="#" class="button" onclick="checkStatus(); return false;">
                📊 Check Migration Status
            </a>
            
            <a href="#" class="button secondary" onclick="startMigration(); return false;">
                🚀 Start Auto-Migration (Tarea1)
            </a>
            
            <a href="#" class="button secondary" onclick="startMigrationTarea2(); return false;">
                🚀 Start Auto-Migration (Tarea2)
            </a>
            
            <a href="#" class="button secondary" onclick="startMigrationTarea3(); return false;">
                🚀 Start Auto-Migration (Tarea3)
            </a>
            
            <a href="#" class="button secondary" onclick="startMigrationTarea4(); return false;">
                🚀 Start Auto-Migration (Tarea4)
            </a>
            
            <a href="#" class="button secondary" onclick="startMigrationTarea5(); return false;">
                🚀 Start Auto-Migration (Tarea5)
            </a>

            <h2>📝 Writing Tasks</h2>
            
            <a href="#" class="button" onclick="checkWritingStatus(); return false;">
                📊 Check Writing Tarea 1 Status
            </a>
            
            <a href="#" class="button secondary" onclick="recreateAndParseWriting(); return false;">
                🔄 Recreate Tables & Parse All Documents
            </a>
            
            <a href="#" class="button secondary" onclick="addBucketUrlColumn(); return false;">
                ➕ Add bucket_url Column (if missing)
            </a>
            
            <a href="#" class="button secondary" onclick="startMigrationWritingTarea1(); return false;">
                🚀 Start Audio Migration (Writing Tarea1)
            </a>

            <h2>Available Endpoints</h2>
            
            <div class="endpoint">
                <strong>GET</strong> /migration-status - Check current migration status
            </div>
            <div class="endpoint">
                <strong>GET</strong> /start-migration - ⭐ Start automatic migration of all tarea1 files
            </div>
            <div class="endpoint">
                <strong>POST</strong> /migrate-audio-files?limit=10 - Migrate specific number of tarea1 files
            </div>
            <div class="endpoint">
                <strong>GET</strong> /start-migration-tarea2 - ⭐ Start automatic migration of all tarea2 files
            </div>
            <div class="endpoint">
                <strong>POST</strong> /migrate-audio-files-tarea2?limit=10 - Migrate specific number of tarea2 files
            </div>
            <div class="endpoint">
                <strong>GET</strong> /start-migration-tarea3 - ⭐ Start automatic migration of all tarea3 files
            </div>
            <div class="endpoint">
                <strong>POST</strong> /migrate-audio-files-tarea3?limit=10 - Migrate specific number of tarea3 files
            </div>
            <div class="endpoint">
                <strong>GET</strong> /start-migration-tarea4 - ⭐ Start automatic migration of all tarea4 files
            </div>
            <div class="endpoint">
                <strong>POST</strong> /migrate-audio-files-tarea4?limit=10 - Migrate specific number of tarea4 files
            </div>
            <div class="endpoint">
                <strong>GET</strong> /start-migration-tarea5 - ⭐ Start automatic migration of all tarea5 files
            </div>
            <div class="endpoint">
                <strong>POST</strong> /migrate-audio-files-tarea5?limit=10 - Migrate specific number of tarea5 files
            </div>
            <div class="endpoint">
                <strong>GET</strong> /audio-files - List all migrated audio files with play URLs
            </div>
            <div class="endpoint">
                <strong>GET</strong> /audio/{file_id} - Get specific audio file by ID
            </div>
            <div class="endpoint">
                <strong>GET</strong> /test-gcs - Test GCS bucket access
            </div>
            <div class="endpoint">
                <strong>GET</strong> /writing-tarea1-status - ⭐ Check Writing Tarea 1 parsing status
            </div>
            <div class="endpoint">
                <strong>POST</strong> /parse-writing-tarea1?limit=10 - Parse Writing Tarea 1 documents
            </div>

            <h2>Audio Player</h2>
            <button class="button" onclick="loadAudioFiles()">
                🎵 Load Audio Files
            </button>
            <button class="button secondary" onclick="loadBucketAudioFiles()">
                ☁️ Load From Bucket (Direct)
            </button>
            
            <div id="audioList"></div>
            <div id="bucketAudioList"></div>

            <div class="loading" id="loading">
                <p>⏳ Processing... Please wait...</p>
            </div>

            <div id="result"></div>
        </div>

        <script>
            function showLoading() {
                document.getElementById('loading').style.display = 'block';
                document.getElementById('result').innerHTML = '';
            }

            function hideLoading() {
                document.getElementById('loading').style.display = 'none';
            }

            function showResult(data, isError = false) {
                hideLoading();
                const resultDiv = document.getElementById('result');
                resultDiv.className = 'result ' + (isError ? 'error' : 'success');
                resultDiv.innerHTML = JSON.stringify(data, null, 2);
            }

            async function checkStatus() {
                showLoading();
                try {
                    const response = await fetch('/migration-status');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigration() {
                if (!confirm('Start automatic migration of all pending TAREA1 files?\\n\\nThis will migrate all files in batches of 10.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigrationTarea2() {
                if (!confirm('Start automatic migration of all pending TAREA2 files?\\n\\nThis will migrate all files in batches of 10.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration-tarea2');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigrationTarea3() {
                if (!confirm('Start automatic migration of all pending TAREA3 files?\\n\\nThis will migrate all files in batches of 10.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration-tarea3');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigrationTarea4() {
                if (!confirm('Start automatic migration of all pending TAREA4 files?\\n\\nThis will migrate all files in batches of 10.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration-tarea4');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigrationTarea5() {
                if (!confirm('Start automatic migration of all pending TAREA5 files?\\n\\nThis will migrate all files in batches of 10.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration-tarea5');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function startMigrationWritingTarea1() {
                if (!confirm('Start automatic migration of all Writing Tarea 1 audio files?\\n\\nThis will migrate audio from Google Drive to GCS bucket.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/start-migration-writing-tarea1');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function addBucketUrlColumn() {
                if (!confirm('Add bucket_url column to writing_tarea1_set table?\\n\\nThis is safe to run multiple times.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/add-bucket-url-column-writing-tarea1', {
                        method: 'POST'
                    });
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function checkWritingStatus() {
                showLoading();
                try {
                    const response = await fetch('/writing-tarea1-status');
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function parseWritingTarea1() {
                if (!confirm('Parse all Writing Tarea 1 documents?\\n\\nThis will extract data from .docx files and insert into database.')) {
                    return;
                }
                
                showLoading();
                try {
                    const response = await fetch('/parse-writing-tarea1', {
                        method: 'POST'
                    });
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function recreateAndParseWriting() {
                if (!confirm('⚠️ WARNING: This will DROP and RECREATE writing_tarea1 tables, then parse ALL 100 documents.\\n\\nAll existing data will be DELETED.\\n\\nContinue?')) {
                    return;
                }
                
                showLoading();
                document.getElementById('result').innerHTML = '<div class="result">🔄 Recreating tables and parsing all documents... This may take a few minutes...</div>';
                
                try {
                    const response = await fetch('/recreate-and-parse-writing-tarea1', {
                        method: 'POST'
                    });
                    const data = await response.json();
                    showResult(data);
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function loadAudioFiles() {
                showLoading();
                try {
                    const response = await fetch('/audio-files?limit=20');
                    const data = await response.json();
                    hideLoading();
                    
                    if (data.success && data.files && data.files.length > 0) {
                        const audioListDiv = document.getElementById('audioList');
                        audioListDiv.innerHTML = '<h3>📁 Audio Files (' + data.total + ' total)</h3>';
                        
                        data.files.forEach((file, index) => {
                            const fileDiv = document.createElement('div');
                            fileDiv.className = 'audio-item';
                            fileDiv.innerHTML = `
                                <div style="background: #f9f9f9; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #4CAF50;">
                                    <strong>File ${index + 1} (ID: ${file.id})</strong><br>
                                    <small>Path: ${file.blob_name}</small><br>
                                    <audio controls style="width: 100%; margin-top: 10px;">
                                        <source src="${file.play_url}" type="audio/mpeg">
                                        Your browser does not support the audio element.
                                    </audio>
                                </div>
                            `;
                            audioListDiv.appendChild(fileDiv);
                        });
                        
                        if (data.total > data.files.length) {
                            audioListDiv.innerHTML += `<p><em>Showing ${data.files.length} of ${data.total} files. Adjust limit parameter to see more.</em></p>`;
                        }
                    } else {
                        showResult(data, !data.success);
                    }
                } catch (error) {
                    showResult({error: error.message}, true);
                }
            }

            async function loadBucketAudioFiles() {
                showLoading();
                try {
                    const response = await fetch('/bucket-audio-files?prefix=listening_tarea1/&limit=20');
                    const data = await response.json();
                    hideLoading();
                    const target = document.getElementById('bucketAudioList');
                    if (!data.success) {
                        target.innerHTML = `<div class='result error'>${JSON.stringify(data, null, 2)}</div>`;
                        return;
                    }
                    target.innerHTML = `<h3>🗂️ Bucket Objects (${data.count}/${data.limit} shown)</h3>`;
                    if (data.files.length === 0) {
                        target.innerHTML += '<p><em>No objects found under given prefix.</em></p>';
                    } else {
                        data.files.forEach((file, idx) => {
                            const div = document.createElement('div');
                            div.className = 'bucket-audio-item';
                            const playable = file.play_url ? `<audio controls style='width:100%;margin-top:8px'><source src='${file.play_url}' type='audio/mpeg'></audio>` : `<p style='color:#d9534f'><em>No signed URL (viewer role only). Use /bucket-audio/${file.blob_name} to stream via backend.</em></p><audio controls style='width:100%;margin-top:8px'><source src='/bucket-audio/${file.blob_name}' type='audio/mpeg'></audio>`;
                            div.innerHTML = `
                                <div style='background:#f9f9f9;padding:12px;margin:8px 0;border-radius:5px;border-left:4px solid #2196F3;'>
                                    <strong>Object ${idx + 1}</strong><br>
                                    <small>${file.blob_name}</small><br>
                                    <small>Size: ${file.size} bytes</small><br>
                                    ${playable}
                                </div>
                            `;
                            target.appendChild(div);
                        });
                    }
                } catch (error) {
                    hideLoading();
                    document.getElementById('bucketAudioList').innerHTML = `<div class='result error'>${error.message}</div>`;
                }
            }

            // Auto-load status on page load
            window.onload = function() {
                checkStatus();
            };
        </script>
    </body>
    </html>
    """
    return html_content

@app.get("/ping")
def ping():
    return {"status": "ok"}

# ----------------------------------------------------------
# 🚀 Test Endpoint to Check GCS Permissions
# ----------------------------------------------------------
@app.get("/test-gcs")
def test_gcs_access():
    """
    Test read/write permissions on the client's GCS bucket
    """
    try:
        client = storage.Client()
        bucket_name = "clarodele-mvp-content"  # Client's bucket
        bucket = client.bucket(bucket_name)

        # Try reading first 2 files
        blobs = list(bucket.list_blobs(max_results=2))
        blob_names = [b.name for b in blobs] if blobs else []

        # Try writing a small test file
        test_blob = bucket.blob("test_access_from_cloudrun.txt")
        test_blob.upload_from_string("✅ Cloud Run has access to this bucket.")

        return {
            "success": True,
            "message": "Read and write access confirmed!",
            "example_files": blob_names,
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ----------------------------------------------------------
# 🔄 Migrate Audio Files from Google Drive to GCS
# ----------------------------------------------------------
@app.post("/migrate-audio-files")
async def migrate_audio_files(limit: Optional[int] = None, test_mode: bool = False):
    """
    Migrate audio files from Google Drive to GCS bucket
    
    Args:
        limit: Number of rows to process (None = all rows)
        test_mode: If True, only process rows without committing to database
    
    Returns:
        Migration results with success/failure counts
    """
    conn = None
    results = {
        "total_rows": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "errors": []
    }
    
    try:
        # Connect to database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # First, check if bucket_url column exists, if not create it
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea1_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            print(f"Column might already exist: {e}")
            conn.rollback()
        
        # Fetch rows that need migration
        query = """
            SELECT tarea1_set_id as id, audio_url, bucket_url 
            FROM listening_tarea1_set 
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL OR bucket_url = '')
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        rows = cursor.fetchall()
        results["total_rows"] = len(rows)
        
        print(f"Found {len(rows)} rows to process")
        
        # Process each row
        for row in rows:
            row_id = row['id']
            audio_url = row['audio_url']
            
            try:
                # Extract Google Drive file ID
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    results["failed"] += 1
                    results["errors"].append({
                        "row_id": row_id,
                        "error": "Could not extract Google Drive file ID",
                        "url": audio_url
                    })
                    continue
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file_path = tmp_file.name
                
                try:
                    # Download from Google Drive
                    print(f"Downloading file {file_id} for row {row_id}...")
                    if not download_from_google_drive(file_id, tmp_file_path):
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to download from Google Drive",
                            "file_id": file_id
                        })
                        continue
                    
                    # Upload to GCS with organized path
                    destination_blob_name = f"listening_tarea1/{file_id}.mp3"
                    print(f"Uploading to GCS: {destination_blob_name}...")
                    bucket_url = upload_to_gcs(tmp_file_path, destination_blob_name)
                    
                    if not bucket_url:
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to upload to GCS",
                            "file_id": file_id
                        })
                        continue
                    
                    # Update database
                    if not test_mode:
                        update_bucket_url(conn, row_id, bucket_url, "listening_tarea1_set", "tarea1_set_id")
                        print(f"Updated row {row_id} with bucket URL: {bucket_url}")
                    else:
                        print(f"[TEST MODE] Would update row {row_id} with: {bucket_url}")
                    
                    results["successful"] += 1
                    
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "row_id": row_id,
                    "error": str(e),
                    "url": audio_url
                })
                print(f"Error processing row {row_id}: {str(e)}")
        
        cursor.close()
        
        return {
            "success": True,
            "message": "Migration completed",
            "results": results,
            "test_mode": test_mode
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": results
        }
    finally:
        if conn:
            conn.close()


@app.get("/start-migration")
async def start_auto_migration():
    """
    🚀 AUTO-MIGRATION ENDPOINT
    
    This endpoint automatically migrates all pending audio files in batches.
    Just visit this URL and it will handle everything!
    
    It will:
    1. Check current migration status
    2. Migrate all pending files in batches of 10
    3. Return final statistics
    
    Safe to run multiple times - skips already migrated files.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Ensure column exists
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea1_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
        
        # Get initial stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea1_set;
        """)
        initial_stats = dict(cursor.fetchone())
        
        cursor.close()
        conn.close()
        
        # If nothing to migrate
        if initial_stats['pending_migration'] == 0:
            return {
                "success": True,
                "message": "✅ All files already migrated!",
                "statistics": initial_stats,
                "migrated_in_this_run": 0
            }
        
        # Migrate all pending files (batch size: 10)
        BATCH_SIZE = 10
        total_migrated = 0
        all_results = []
        
        while True:
            # Migrate one batch
            result = await migrate_audio_files(limit=BATCH_SIZE, test_mode=False)
            
            if not result.get("success"):
                return {
                    "success": False,
                    "error": "Migration failed",
                    "details": result,
                    "total_migrated_before_error": total_migrated
                }
            
            batch_success = result["results"]["successful"]
            total_migrated += batch_success
            all_results.append(result["results"])
            
            # If this batch had no successful migrations, we're done
            if batch_success == 0:
                break
        
        # Get final stats
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea1_set;
        """)
        final_stats = dict(cursor.fetchone())
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"✅ Migration completed! Migrated {total_migrated} files.",
            "migrated_in_this_run": total_migrated,
            "initial_statistics": initial_stats,
            "final_statistics": final_stats,
            "batch_results": all_results
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


# ----------------------------------------------------------
# 🔄 Migrate Audio Files for Tarea2 Table
# ----------------------------------------------------------
@app.post("/migrate-audio-files-tarea2")
async def migrate_audio_files_tarea2(limit: Optional[int] = None, test_mode: bool = False):
    """
    Migrate audio files from Google Drive to GCS bucket for listening_tarea2_set table
    
    Args:
        limit: Number of rows to process (None = all rows)
        test_mode: If True, only process rows without committing to database
    
    Returns:
        Migration results with success/failure counts
    """
    conn = None
    results = {
        "total_rows": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "errors": []
    }
    
    try:
        # Connect to database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # First, check if bucket_url column exists, if not create it
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea2_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            print(f"Column might already exist: {e}")
            conn.rollback()
        
        # Fetch rows that need migration
        query = """
            SELECT tarea2_set_id as id, audio_url, bucket_url 
            FROM listening_tarea2_set 
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL OR bucket_url = '')
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        rows = cursor.fetchall()
        results["total_rows"] = len(rows)
        
        print(f"Found {len(rows)} rows to process in listening_tarea2_set")
        
        # Process each row
        for row in rows:
            row_id = row['id']
            audio_url = row['audio_url']
            
            try:
                # Extract Google Drive file ID
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    results["failed"] += 1
                    results["errors"].append({
                        "row_id": row_id,
                        "error": "Could not extract Google Drive file ID",
                        "url": audio_url
                    })
                    continue
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file_path = tmp_file.name
                
                try:
                    # Download from Google Drive
                    print(f"Downloading file {file_id} for tarea2 row {row_id}...")
                    if not download_from_google_drive(file_id, tmp_file_path):
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to download from Google Drive",
                            "file_id": file_id
                        })
                        continue
                    
                    # Upload to GCS with organized path
                    destination_blob_name = f"listening_tarea2/{file_id}.mp3"
                    print(f"Uploading to GCS: {destination_blob_name}...")
                    bucket_url = upload_to_gcs(tmp_file_path, destination_blob_name)
                    
                    if not bucket_url:
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to upload to GCS",
                            "file_id": file_id
                        })
                        continue
                    
                    # Update database
                    if not test_mode:
                        update_bucket_url(conn, row_id, bucket_url, "listening_tarea2_set", "tarea2_set_id")
                        print(f"Updated tarea2 row {row_id} with bucket URL: {bucket_url}")
                    else:
                        print(f"[TEST MODE] Would update tarea2 row {row_id} with: {bucket_url}")
                    
                    results["successful"] += 1
                    
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "row_id": row_id,
                    "error": str(e),
                    "url": audio_url
                })
                print(f"Error processing tarea2 row {row_id}: {str(e)}")
        
        cursor.close()
        
        return {
            "success": True,
            "message": "Migration completed for listening_tarea2_set",
            "results": results,
            "test_mode": test_mode
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": results
        }
    finally:
        if conn:
            conn.close()


@app.get("/start-migration-tarea2")
async def start_auto_migration_tarea2():
    """
    🚀 AUTO-MIGRATION ENDPOINT FOR TAREA2
    
    This endpoint automatically migrates all pending audio files in batches.
    Just visit this URL and it will handle everything!
    
    It will:
    1. Check current migration status
    2. Migrate all pending files in batches of 10
    3. Return final statistics
    
    Safe to run multiple times - skips already migrated files.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Ensure column exists
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea2_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
        
        # Get initial stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea2_set;
        """)
        initial_stats = dict(cursor.fetchone())
        
        cursor.close()
        conn.close()
        
        # If nothing to migrate
        if initial_stats['pending_migration'] == 0:
            return {
                "success": True,
                "message": "✅ All tarea2 files already migrated!",
                "statistics": initial_stats,
                "migrated_in_this_run": 0
            }
        
        # Migrate all pending files (batch size: 10)
        BATCH_SIZE = 10
        total_migrated = 0
        all_results = []
        
        while True:
            # Migrate one batch
            result = await migrate_audio_files_tarea2(limit=BATCH_SIZE, test_mode=False)
            
            if not result.get("success"):
                return {
                    "success": False,
                    "error": "Migration failed",
                    "details": result,
                    "total_migrated_before_error": total_migrated
                }
            
            batch_success = result["results"]["successful"]
            total_migrated += batch_success
            all_results.append(result["results"])
            
            # If this batch had no successful migrations, we're done
            if batch_success == 0:
                break
        
        # Get final stats
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea2_set;
        """)
        final_stats = dict(cursor.fetchone())
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"✅ Tarea2 migration completed! Migrated {total_migrated} files.",
            "migrated_in_this_run": total_migrated,
            "initial_statistics": initial_stats,
            "final_statistics": final_stats,
            "batch_results": all_results
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


# ----------------------------------------------------------
# 🔄 Migrate Audio Files for Tarea3 Table
# ----------------------------------------------------------
@app.post("/migrate-audio-files-tarea3")
async def migrate_audio_files_tarea3(limit: Optional[int] = None, test_mode: bool = False):
    """
    Migrate audio files from Google Drive to GCS bucket for listening_tarea3_set table
    
    Args:
        limit: Number of rows to process (None = all rows)
        test_mode: If True, only process rows without committing to database
    
    Returns:
        Migration results with success/failure counts
    """
    conn = None
    results = {
        "total_rows": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "errors": []
    }
    
    try:
        # Connect to database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # First, check if bucket_url column exists, if not create it
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea3_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            print(f"Column might already exist: {e}")
            conn.rollback()
        
        # Fetch rows that need migration
        query = """
            SELECT tarea3_set_id as id, audio_url, bucket_url 
            FROM listening_tarea3_set 
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL OR bucket_url = '')
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        rows = cursor.fetchall()
        results["total_rows"] = len(rows)
        
        print(f"Found {len(rows)} rows to process in listening_tarea3_set")
        
        # Process each row
        for row in rows:
            row_id = row['id']
            audio_url = row['audio_url']
            
            try:
                # Extract Google Drive file ID
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    results["failed"] += 1
                    results["errors"].append({
                        "row_id": row_id,
                        "error": "Could not extract Google Drive file ID",
                        "url": audio_url
                    })
                    continue
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file_path = tmp_file.name
                
                try:
                    # Download from Google Drive
                    print(f"Downloading file {file_id} for tarea3 row {row_id}...")
                    if not download_from_google_drive(file_id, tmp_file_path):
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to download from Google Drive",
                            "file_id": file_id
                        })
                        continue
                    
                    # Upload to GCS with organized path
                    destination_blob_name = f"listening_tarea3/{file_id}.mp3"
                    print(f"Uploading to GCS: {destination_blob_name}...")
                    bucket_url = upload_to_gcs(tmp_file_path, destination_blob_name)
                    
                    if not bucket_url:
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to upload to GCS",
                            "file_id": file_id
                        })
                        continue
                    
                    # Update database
                    if not test_mode:
                        update_bucket_url(conn, row_id, bucket_url, "listening_tarea3_set", "tarea3_set_id")
                        print(f"Updated tarea3 row {row_id} with bucket URL: {bucket_url}")
                    else:
                        print(f"[TEST MODE] Would update tarea3 row {row_id} with: {bucket_url}")
                    
                    results["successful"] += 1
                    
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "row_id": row_id,
                    "error": str(e),
                    "url": audio_url
                })
                print(f"Error processing tarea3 row {row_id}: {str(e)}")
        
        cursor.close()
        
        return {
            "success": True,
            "message": "Migration completed for listening_tarea3_set",
            "results": results,
            "test_mode": test_mode
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": results
        }
    finally:
        if conn:
            conn.close()


@app.get("/start-migration-tarea3")
async def start_auto_migration_tarea3():
    """
    🚀 AUTO-MIGRATION ENDPOINT FOR TAREA3
    
    This endpoint automatically migrates all pending audio files in batches.
    Just visit this URL and it will handle everything!
    
    It will:
    1. Check current migration status
    2. Migrate all pending files in batches of 10
    3. Return final statistics
    
    Safe to run multiple times - skips already migrated files.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Ensure column exists
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea3_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
        
        # Get initial stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea3_set;
        """)
        initial_stats = dict(cursor.fetchone())
        
        cursor.close()
        conn.close()
        
        # If nothing to migrate
        if initial_stats['pending_migration'] == 0:
            return {
                "success": True,
                "message": "✅ All tarea3 files already migrated!",
                "statistics": initial_stats,
                "migrated_in_this_run": 0
            }
        
        # Migrate all pending files (batch size: 10)
        BATCH_SIZE = 10
        total_migrated = 0
        all_results = []
        
        while True:
            # Migrate one batch
            result = await migrate_audio_files_tarea3(limit=BATCH_SIZE, test_mode=False)
            
            if not result.get("success"):
                return {
                    "success": False,
                    "error": "Migration failed",
                    "details": result,
                    "total_migrated_before_error": total_migrated
                }
            
            batch_success = result["results"]["successful"]
            total_migrated += batch_success
            all_results.append(result["results"])
            
            # If this batch had no successful migrations, we're done
            if batch_success == 0:
                break
        
        # Get final stats
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea3_set;
        """)
        final_stats = dict(cursor.fetchone())
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"✅ Tarea3 migration completed! Migrated {total_migrated} files.",
            "migrated_in_this_run": total_migrated,
            "initial_statistics": initial_stats,
            "final_statistics": final_stats,
            "batch_results": all_results
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/migrate-audio-files-tarea4")
async def migrate_audio_files_tarea4(limit: Optional[int] = None, test_mode: bool = False):
    """
    Migrate audio files from Google Drive to GCS bucket for listening_tarea4_set table
    
    Args:
        limit: Number of rows to process (None = all rows)
        test_mode: If True, only process rows without committing to database
    
    Returns:
        Migration results with success/failure counts
    """
    conn = None
    results = {
        "total_rows": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "errors": []
    }
    
    try:
        # Connect to database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # First, check if bucket_url column exists, if not create it
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea4_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            print(f"Column might already exist: {e}")
            conn.rollback()
        
        # Fetch rows that need migration
        query = """
            SELECT tarea4_set_id as id, audio_url, bucket_url 
            FROM listening_tarea4_set 
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL OR bucket_url = '')
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        rows = cursor.fetchall()
        results["total_rows"] = len(rows)
        
        print(f"Found {len(rows)} rows to process in listening_tarea4_set")
        
        # Process each row
        for row in rows:
            row_id = row['id']
            audio_url = row['audio_url']
            
            try:
                # Extract Google Drive file ID
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    results["failed"] += 1
                    results["errors"].append({
                        "row_id": row_id,
                        "error": "Could not extract Google Drive file ID",
                        "url": audio_url
                    })
                    continue
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file_path = tmp_file.name
                
                try:
                    # Download from Google Drive
                    print(f"Downloading file {file_id} for tarea4 row {row_id}...")
                    if not download_from_google_drive(file_id, tmp_file_path):
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to download from Google Drive",
                            "file_id": file_id
                        })
                        continue
                    
                    # Upload to GCS with organized path
                    destination_blob_name = f"listening_tarea4/{file_id}.mp3"
                    print(f"Uploading to GCS: {destination_blob_name}...")
                    bucket_url = upload_to_gcs(tmp_file_path, destination_blob_name)
                    
                    if not bucket_url:
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to upload to GCS",
                            "file_id": file_id
                        })
                        continue
                    
                    # Update database
                    if not test_mode:
                        update_bucket_url(conn, row_id, bucket_url, "listening_tarea4_set", "tarea4_set_id")
                        print(f"Updated tarea4 row {row_id} with bucket URL: {bucket_url}")
                    else:
                        print(f"[TEST MODE] Would update tarea4 row {row_id} with: {bucket_url}")
                    
                    results["successful"] += 1
                    
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "row_id": row_id,
                    "error": str(e),
                    "url": audio_url
                })
                print(f"Error processing tarea4 row {row_id}: {str(e)}")
        
        cursor.close()
        
        return {
            "success": True,
            "message": "Migration completed for listening_tarea4_set",
            "results": results,
            "test_mode": test_mode
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": results
        }
    finally:
        if conn:
            conn.close()


@app.get("/start-migration-tarea4")
async def start_auto_migration_tarea4():
    """
    🚀 AUTO-MIGRATION ENDPOINT FOR TAREA4
    
    This endpoint automatically migrates all pending audio files in batches.
    Just visit this URL and it will handle everything!
    
    It will:
    1. Check current migration status
    2. Migrate all pending files in batches of 10
    3. Return final statistics
    
    Safe to run multiple times - skips already migrated files.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Ensure column exists
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea4_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
        
        # Get initial stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea4_set;
        """)
        initial_stats = dict(cursor.fetchone())
        
        cursor.close()
        conn.close()
        
        # If nothing to migrate
        if initial_stats['pending_migration'] == 0:
            return {
                "success": True,
                "message": "✅ All tarea4 files already migrated!",
                "statistics": initial_stats,
                "migrated_in_this_run": 0
            }
        
        # Migrate all pending files (batch size: 10)
        BATCH_SIZE = 10
        total_migrated = 0
        all_results = []
        
        while True:
            # Migrate one batch
            result = await migrate_audio_files_tarea4(limit=BATCH_SIZE, test_mode=False)
            
            if not result.get("success"):
                return {
                    "success": False,
                    "error": "Migration failed",
                    "details": result,
                    "total_migrated_before_error": total_migrated
                }
            
            batch_success = result["results"]["successful"]
            total_migrated += batch_success
            all_results.append(result["results"])
            
            # If this batch had no successful migrations, we're done
            if batch_success == 0:
                break
        
        # Get final stats
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(bucket_url) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
            FROM listening_tarea4_set;
        """)
        final_stats = dict(cursor.fetchone())
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"✅ Tarea4 migration completed! Migrated {total_migrated} files.",
            "migrated_in_this_run": total_migrated,
            "initial_statistics": initial_stats,
            "final_statistics": final_stats,
            "batch_results": all_results
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/migrate-audio-files-tarea5")
async def migrate_audio_files_tarea5(limit: Optional[int] = None, test_mode: bool = False):
    """
    Migrate audio files from Google Drive to GCS bucket for listening_tarea5_set table
    
    Args:
        limit: Number of rows to process (None = all rows)
        test_mode: If True, only process rows without committing to database
    
    Returns:
        Migration results with success/failure counts
    """
    conn = None
    results = {
        "total_rows": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "errors": []
    }
    
    try:
        # Connect to database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # First, check if bucket_url column exists, if not create it
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea5_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            print(f"Column might already exist: {e}")
            conn.rollback()
        
        # Fetch rows that need migration
        # Include rows where bucket_url is NULL, empty, or contains Google Drive links
        query = """
            SELECT tarea5_set_id as id, audio_url, bucket_url 
            FROM listening_tarea5_set 
            WHERE audio_url IS NOT NULL 
            AND (
                bucket_url IS NULL 
                OR bucket_url = '' 
                OR bucket_url LIKE '%drive.google.com%'
                OR bucket_url LIKE '%docs.google.com%'
                OR bucket_url NOT LIKE 'gs://%'
            )
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        rows = cursor.fetchall()
        results["total_rows"] = len(rows)
        
        print(f"Found {len(rows)} rows to process in listening_tarea5_set")
        
        # Process each row
        for row in rows:
            row_id = row['id']
            audio_url = row['audio_url']
            
            try:
                # Extract Google Drive file ID
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    results["failed"] += 1
                    results["errors"].append({
                        "row_id": row_id,
                        "error": "Could not extract Google Drive file ID",
                        "url": audio_url
                    })
                    continue
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                    tmp_file_path = tmp_file.name
                
                try:
                    # Download from Google Drive
                    print(f"Downloading file {file_id} for tarea5 row {row_id}...")
                    if not download_from_google_drive(file_id, tmp_file_path):
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to download from Google Drive",
                            "file_id": file_id
                        })
                        continue
                    
                    # Upload to GCS with organized path
                    destination_blob_name = f"listening_tarea5/{file_id}.mp3"
                    print(f"Uploading to GCS: {destination_blob_name}...")
                    bucket_url = upload_to_gcs(tmp_file_path, destination_blob_name)
                    
                    if not bucket_url:
                        results["failed"] += 1
                        results["errors"].append({
                            "row_id": row_id,
                            "error": "Failed to upload to GCS",
                            "file_id": file_id
                        })
                        continue
                    
                    # Update database
                    if not test_mode:
                        update_bucket_url(conn, row_id, bucket_url, "listening_tarea5_set", "tarea5_set_id")
                        print(f"Updated tarea5 row {row_id} with bucket URL: {bucket_url}")
                    else:
                        print(f"[TEST MODE] Would update tarea5 row {row_id} with: {bucket_url}")
                    
                    results["successful"] += 1
                    
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "row_id": row_id,
                    "error": str(e),
                    "url": audio_url
                })
                print(f"Error processing tarea5 row {row_id}: {str(e)}")
        
        cursor.close()
        
        return {
            "success": True,
            "message": "Migration completed for listening_tarea5_set",
            "results": results,
            "test_mode": test_mode
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": results
        }
    finally:
        if conn:
            conn.close()


@app.get("/start-migration-tarea5")
async def start_auto_migration_tarea5():
    """
    🚀 AUTO-MIGRATION ENDPOINT FOR TAREA5
    
    This endpoint automatically migrates all pending audio files in batches.
    Just visit this URL and it will handle everything!
    
    It will:
    1. Check current migration status
    2. Migrate all pending files in batches of 10
    3. Return final statistics
    
    Safe to run multiple times - skips already migrated files.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Ensure column exists
        try:
            cursor.execute("""
                ALTER TABLE listening_tarea5_set 
                ADD COLUMN IF NOT EXISTS bucket_url TEXT;
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
        
        # Get initial stats
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(CASE WHEN bucket_url IS NOT NULL AND bucket_url LIKE 'gs://%' THEN 1 END) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (
                    bucket_url IS NULL 
                    OR bucket_url = '' 
                    OR bucket_url LIKE '%drive.google.com%'
                    OR bucket_url LIKE '%docs.google.com%'
                    OR bucket_url NOT LIKE 'gs://%'
                ) THEN 1 END) as pending_migration
            FROM listening_tarea5_set;
        """)
        initial_stats = dict(cursor.fetchone())
        
        cursor.close()
        conn.close()
        
        # If nothing to migrate
        if initial_stats['pending_migration'] == 0:
            return {
                "success": True,
                "message": "✅ All tarea5 files already migrated!",
                "statistics": initial_stats,
                "migrated_in_this_run": 0
            }
        
        # Migrate all pending files (batch size: 10)
        BATCH_SIZE = 10
        total_migrated = 0
        all_results = []
        
        while True:
            # Migrate one batch
            result = await migrate_audio_files_tarea5(limit=BATCH_SIZE, test_mode=False)
            
            if not result.get("success"):
                return {
                    "success": False,
                    "error": "Migration failed",
                    "details": result,
                    "total_migrated_before_error": total_migrated
                }
            
            batch_success = result["results"]["successful"]
            total_migrated += batch_success
            all_results.append(result["results"])
            
            # If this batch had no successful migrations, we're done
            if batch_success == 0:
                break
        
        # Get final stats
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute("""
            SELECT 
                COUNT(*) as total_rows,
                COUNT(audio_url) as rows_with_audio_url,
                COUNT(CASE WHEN bucket_url IS NOT NULL AND bucket_url LIKE 'gs://%' THEN 1 END) as rows_with_bucket_url,
                COUNT(CASE WHEN audio_url IS NOT NULL AND (
                    bucket_url IS NULL 
                    OR bucket_url = '' 
                    OR bucket_url LIKE '%drive.google.com%'
                    OR bucket_url LIKE '%docs.google.com%'
                    OR bucket_url NOT LIKE 'gs://%'
                ) THEN 1 END) as pending_migration
            FROM listening_tarea5_set;
        """)
        final_stats = dict(cursor.fetchone())
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"✅ Tarea5 migration completed! Migrated {total_migrated} files.",
            "migrated_in_this_run": total_migrated,
            "initial_statistics": initial_stats,
            "final_statistics": final_stats,
            "batch_results": all_results
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/migration-status")
async def get_migration_status():
    """
    Check the migration status of audio files for all tarea tables (listening + writing)
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Check tarea1 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='listening_tarea1_set' 
            AND column_name='bucket_url';
        """)
        tarea1_column_exists = cursor.fetchone() is not None
        
        # Check tarea2 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='listening_tarea2_set' 
            AND column_name='bucket_url';
        """)
        tarea2_column_exists = cursor.fetchone() is not None
        
        # Check tarea3 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='listening_tarea3_set' 
            AND column_name='bucket_url';
        """)
        tarea3_column_exists = cursor.fetchone() is not None
        
        # Check tarea4 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='listening_tarea4_set' 
            AND column_name='bucket_url';
        """)
        tarea4_column_exists = cursor.fetchone() is not None
        
        # Check tarea5 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='listening_tarea5_set' 
            AND column_name='bucket_url';
        """)
        tarea5_column_exists = cursor.fetchone() is not None
        
        # Check writing_tarea1 table
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='writing_tarea1_set' 
            AND column_name='bucket_url';
        """)
        writing_tarea1_column_exists = cursor.fetchone() is not None
        
        result = {
            "success": True,
            "tarea1": {},
            "tarea2": {},
            "tarea3": {},
            "tarea4": {},
            "tarea5": {},
            "writing_tarea1": {}
        }
        
        # Get tarea1 statistics
        if tarea1_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(bucket_url) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
                FROM listening_tarea1_set;
            """)
            tarea1_stats = dict(cursor.fetchone())
            result["tarea1"] = {
                "column_exists": True,
                "statistics": tarea1_stats
            }
        else:
            result["tarea1"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        # Get tarea2 statistics
        if tarea2_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(bucket_url) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
                FROM listening_tarea2_set;
            """)
            tarea2_stats = dict(cursor.fetchone())
            result["tarea2"] = {
                "column_exists": True,
                "statistics": tarea2_stats
            }
        else:
            result["tarea2"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        # Get tarea3 statistics
        if tarea3_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(bucket_url) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
                FROM listening_tarea3_set;
            """)
            tarea3_stats = dict(cursor.fetchone())
            result["tarea3"] = {
                "column_exists": True,
                "statistics": tarea3_stats
            }
        else:
            result["tarea3"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        # Get tarea4 statistics
        if tarea4_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(bucket_url) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (bucket_url IS NULL OR bucket_url = '') THEN 1 END) as pending_migration
                FROM listening_tarea4_set;
            """)
            tarea4_stats = dict(cursor.fetchone())
            result["tarea4"] = {
                "column_exists": True,
                "statistics": tarea4_stats
            }
        else:
            result["tarea4"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        # Get tarea5 statistics
        if tarea5_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(CASE WHEN bucket_url IS NOT NULL AND bucket_url LIKE 'gs://%' THEN 1 END) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (
                        bucket_url IS NULL 
                        OR bucket_url = '' 
                        OR bucket_url LIKE '%drive.google.com%'
                        OR bucket_url LIKE '%docs.google.com%'
                        OR bucket_url NOT LIKE 'gs://%'
                    ) THEN 1 END) as pending_migration
                FROM listening_tarea5_set;
            """)
            tarea5_stats = dict(cursor.fetchone())
            result["tarea5"] = {
                "column_exists": True,
                "statistics": tarea5_stats
            }
        else:
            result["tarea5"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        # Get writing_tarea1 statistics
        if writing_tarea1_column_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_rows,
                    COUNT(audio_url) as rows_with_audio_url,
                    COUNT(CASE WHEN bucket_url IS NOT NULL AND bucket_url LIKE 'gs://%' THEN 1 END) as rows_with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND (
                        bucket_url IS NULL 
                        OR bucket_url = '' 
                        OR bucket_url LIKE '%drive.google.com%'
                        OR bucket_url NOT LIKE 'gs://%'
                    ) THEN 1 END) as pending_migration
                FROM writing_tarea1_set;
            """)
            writing_tarea1_stats = dict(cursor.fetchone())
            result["writing_tarea1"] = {
                "column_exists": True,
                "statistics": writing_tarea1_stats
            }
        else:
            result["writing_tarea1"] = {
                "column_exists": False,
                "message": "bucket_url column does not exist yet"
            }
        
        cursor.close()
        conn.close()
        
        return result
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/audio-files")
async def list_audio_files(limit: int = 50, offset: int = 0):
    """
    List all migrated audio files with playable URLs
    
    Args:
        limit: Number of files to return (default: 50)
        offset: Number of files to skip (default: 0)
    
    Returns:
        List of audio files with signed URLs for playing
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get audio files with bucket URLs
        cursor.execute("""
            SELECT 
                tarea1_set_id as id,
                audio_url,
                bucket_url
            FROM listening_tarea1_set
            WHERE bucket_url IS NOT NULL 
            AND bucket_url != ''
            ORDER BY tarea1_set_id
            LIMIT %s OFFSET %s;
        """, (limit, offset))
        
        files = cursor.fetchall()
        
        # Get total count
        cursor.execute("""
            SELECT COUNT(*) as total
            FROM listening_tarea1_set
            WHERE bucket_url IS NOT NULL 
            AND bucket_url != '';
        """)
        total = cursor.fetchone()['total']
        
        cursor.close()
        conn.close()
        
        # Generate signed URLs for each file
        client = storage.Client()
        bucket = client.bucket(GCS_BUCKET_NAME)
        
        audio_files = []
        for file in files:
            bucket_url = file['bucket_url']
            
            # Extract blob name from gs:// URL
            if bucket_url.startswith(f"gs://{GCS_BUCKET_NAME}/"):
                blob_name = bucket_url.replace(f"gs://{GCS_BUCKET_NAME}/", "")
                blob = bucket.blob(blob_name)
                
                # Generate signed URL (valid for 1 hour)
                signed_url = blob.generate_signed_url(
                    version="v4",
                    expiration=3600,  # 1 hour
                    method="GET"
                )
                
                audio_files.append({
                    "id": file['id'],
                    "google_drive_url": file['audio_url'],
                    "bucket_url": bucket_url,
                    "play_url": signed_url,
                    "blob_name": blob_name
                })
        
        return {
            "success": True,
            "total": total,
            "limit": limit,
            "offset": offset,
            "count": len(audio_files),
            "files": audio_files
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/audio/{file_id}")
async def get_audio_file(file_id: int):
    """
    Get a specific audio file by ID with a playable URL
    
    Args:
        file_id: The tarea1_set_id of the audio file
    
    Returns:
        Audio file details with signed URL for playing
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get the specific audio file
        cursor.execute("""
            SELECT 
                tarea1_set_id as id,
                audio_url,
                bucket_url
            FROM listening_tarea1_set
            WHERE tarea1_set_id = %s
            AND bucket_url IS NOT NULL 
            AND bucket_url != '';
        """, (file_id,))
        
        file = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not file:
            return {
                "success": False,
                "error": f"Audio file with ID {file_id} not found or not migrated yet"
            }
        
        bucket_url = file['bucket_url']
        
        # Extract blob name from gs:// URL
        if bucket_url.startswith(f"gs://{GCS_BUCKET_NAME}/"):
            blob_name = bucket_url.replace(f"gs://{GCS_BUCKET_NAME}/", "")
            
            # Generate signed URL
            client = storage.Client()
            bucket = client.bucket(GCS_BUCKET_NAME)
            blob = bucket.blob(blob_name)
            
            signed_url = blob.generate_signed_url(
                version="v4",
                expiration=3600,  # 1 hour
                method="GET"
            )
            
            return {
                "success": True,
                "file": {
                    "id": file['id'],
                    "google_drive_url": file['audio_url'],
                    "bucket_url": bucket_url,
                    "play_url": signed_url,
                    "blob_name": blob_name
                }
            }
        else:
            return {
                "success": False,
                "error": "Invalid bucket URL format"
            }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


# ----------------------------------------------------------
# 📂 List & Stream Audio Directly From Bucket (No DB Needed)
# ----------------------------------------------------------
@app.get("/bucket-audio-files")
async def list_bucket_audio_files(prefix: str = "", limit: int = 50):
    """
    List audio objects directly from the GCS bucket.

    Args:
        prefix: Optional path prefix to filter objects (e.g. 'listening_tarea1/')
        limit: Max number of objects to return

    Returns:
        Metadata for objects plus signed URL if credentials allow signing.

    Notes:
        - Requires at least storage.objects.list permission.
        - Signed URL generation needs a service account with signBlob ability. If
          you only have *viewer* rights via a user credential, signed URL may fail.
        - If signing fails, play_url will be omitted; consider making objects public
          or using a service account key JSON.
    """
    try:
        client = storage.Client()
        bucket = client.bucket(GCS_BUCKET_NAME)

        blobs_iter = bucket.list_blobs(prefix=prefix)
        items = []
        count = 0
        for blob in blobs_iter:
            if count >= limit:
                break
            # Attempt to generate signed URL (best effort)
            play_url = None
            try:
                play_url = blob.generate_signed_url(version="v4", expiration=3600, method="GET")
            except Exception as signing_err:
                # Silently ignore signing error; user may only have viewer role
                play_url = None

            items.append({
                "blob_name": blob.name,
                "size": blob.size,
                "updated": blob.updated.isoformat() if blob.updated else None,
                "content_type": blob.content_type,
                "crc32c": blob.crc32c,
                "md5_hash": blob.md5_hash,
                "storage_class": blob.storage_class,
                "gcs_uri": f"gs://{GCS_BUCKET_NAME}/{blob.name}",
                "play_url": play_url
            })
            count += 1

        return {
            "success": True,
            "bucket": GCS_BUCKET_NAME,
            "prefix": prefix,
            "limit": limit,
            "count": len(items),
            "files": items
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/bucket-audio/{blob_path:path}")
async def stream_bucket_audio(blob_path: str):
    """
    Stream an audio object directly from GCS without relying on DB.

    Args:
        blob_path: Path of the object inside the bucket (e.g. 'listening_tarea1/123.mp3')

    Returns:
        StreamingResponse with audio content.

    Permissions:
        - Requires storage.objects.get.
        - Efficient for small/medium files (< ~10MB). For very large files you
          might prefer signed URLs to let the client download directly.
    """
    try:
        client = storage.Client()
        bucket = client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(blob_path)

        if not blob.exists():
            raise HTTPException(status_code=404, detail="Blob not found in bucket")

        # Download as bytes (could stream in chunks if very large)
        data = blob.download_as_bytes()
        content_type = blob.content_type or "audio/mpeg"
        return StreamingResponse(iter([data]), media_type=content_type)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ----------------------------------------------------------
# 📝 WRITING TAREA 1 PARSING ENDPOINTS
# ----------------------------------------------------------

@app.post("/parse-writing-tarea1")
async def parse_writing_tarea1_endpoint(
    limit: Optional[int] = None,
    module_type_id: int = 1,
    dry_run: bool = False
):
    """
    Parse Writing Tarea 1 documents and insert into database.
    
    Args:
        limit: Number of documents to process (None = all)
        module_type_id: Module type ID for the documents
        dry_run: If True, parse but don't insert into database
    
    Returns:
        Parsing results with success/failure counts
    """
    try:
        from docx import Document
        
        folder_path = "Writing_Tarea_1"
        
        if not os.path.exists(folder_path):
            return {
                "success": False,
                "error": f"Folder '{folder_path}' not found"
            }
        
        # Get all .docx files
        files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
        files.sort()
        
        if limit:
            files = files[:limit]
        
        conn = None if dry_run else get_db_connection()
        
        results = {
            'total': len(files),
            'successful': 0,
            'failed': 0,
            'errors': [],
            'processed_files': []
        }
        
        try:
            for idx, filename in enumerate(files, 1):
                file_path = os.path.join(folder_path, filename)
                
                try:
                    # Parse document
                    doc = Document(file_path)
                    full_text = []
                    
                    # Extract all text from paragraphs
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            full_text.append(text)
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text = cell.text.strip()
                                if text:
                                    full_text.append(text)
                    
                    raw_text = '\n'.join(full_text)
                    
                    # Parse structured data
                    data = parse_writing_document(raw_text, filename)
                    
                    if dry_run:
                        results['processed_files'].append({
                            'file': filename,
                            'title': data.get('title', 'N/A')[:50],
                            'situation_chars': len(data.get('situation', '')),
                            'task_chars': len(data.get('task_instructions', '')),
                            'solution_chars': len(data.get('solution_text', '')),
                            'total_chars': len(raw_text)
                        })
                    else:
                        # Insert into database
                        tarea1_set_id = insert_writing_tarea1_db(conn, data, module_type_id, filename)
                        results['processed_files'].append({
                            'file': filename,
                            'id': tarea1_set_id,
                            'title': data.get('title', 'N/A')[:50]
                        })
                    
                    results['successful'] += 1
                    
                except Exception as e:
                    results['failed'] += 1
                    results['errors'].append({
                        'file': filename,
                        'error': str(e)
                    })
        
        finally:
            if conn:
                conn.close()
        
        return {
            "success": True,
            "message": f"Processed {results['successful']} of {results['total']} files",
            "results": results,
            "dry_run": dry_run
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


def parse_writing_document(text: str, filename: str) -> dict:
    """Parse writing document text into structured data"""
    lines = text.split('\n')
    
    data = {
        "title": None,
        "situation": None,
        "task_instructions": None,
        "word_limit": None,
        "text_type": None,
        "register": None,
        "reminders": None,
        "audio_url": None,
        "solution_text": None
    }
    
    current_section = None
    section_content = []
    
    situation_keywords = ['SITUACIÓN', 'Situación', 'SITUACION']
    task_keywords = ['TAREA', 'Tarea', 'INSTRUCCIONES']
    solution_keywords = ['SOLUCIÓN', 'Solución', 'SOLUCION', 'Modelo de respuesta', 'MODELO']
    reminder_keywords = ['Recuerde', 'RECUERDE', 'Recordatorio']
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        if not line_stripped:
            continue
        
        # Detect sections
        if any(keyword in line_stripped for keyword in situation_keywords):
            if current_section and section_content:
                save_writing_section(data, current_section, section_content)
            current_section = 'situation'
            section_content = []
            if ':' in line_stripped:
                content = line_stripped.split(':', 1)[1].strip()
                if content:
                    section_content.append(content)
            continue
        
        if any(keyword in line_stripped for keyword in task_keywords):
            if current_section and section_content:
                save_writing_section(data, current_section, section_content)
            current_section = 'task'
            section_content = []
            if ':' in line_stripped:
                content = line_stripped.split(':', 1)[1].strip()
                if content:
                    section_content.append(content)
            continue
        
        if any(keyword in line_stripped for keyword in solution_keywords):
            if current_section and section_content:
                save_writing_section(data, current_section, section_content)
            current_section = 'solution'
            section_content = []
            if ':' in line_stripped:
                content = line_stripped.split(':', 1)[1].strip()
                if content:
                    section_content.append(content)
            continue
        
        if any(keyword in line_stripped for keyword in reminder_keywords):
            if current_section and section_content:
                save_writing_section(data, current_section, section_content)
            current_section = 'reminder'
            section_content = []
            if ':' in line_stripped:
                content = line_stripped.split(':', 1)[1].strip()
                if content:
                    section_content.append(content)
            continue
        
        # Extract metadata
        if 'palabras' in line_stripped.lower() and ('150' in line_stripped or '180' in line_stripped):
            data['word_limit'] = line_stripped
            continue
        
        if 'tipo de texto' in line_stripped.lower() or 'tipo:' in line_stripped.lower():
            if ':' in line_stripped:
                data['text_type'] = line_stripped.split(':', 1)[1].strip()
            continue
        
        if 'registro' in line_stripped.lower() and len(line_stripped) < 50:
            if ':' in line_stripped:
                data['register'] = line_stripped.split(':', 1)[1].strip()
            continue
        
        # Title detection
        if data['title'] is None and len(line_stripped) > 5 and len(line_stripped) < 100:
            if i < 3 and not any(k in line_stripped for k in situation_keywords + task_keywords):
                data['title'] = line_stripped
                continue
        
        # Add to current section
        if current_section:
            section_content.append(line_stripped)
    
    # Save last section
    if current_section and section_content:
        save_writing_section(data, current_section, section_content)
    
    # If no structure found, store all in situation
    if not data['situation'] and not data['task_instructions']:
        data['situation'] = text
    
    return data


def save_writing_section(data: dict, section_name: str, content: list):
    """Save section content to data dictionary"""
    text = '\n'.join(content).strip()
    
    if section_name == 'situation':
        data['situation'] = text
    elif section_name == 'task':
        data['task_instructions'] = text
    elif section_name == 'solution':
        data['solution_text'] = text
    elif section_name == 'reminder':
        data['reminders'] = text


def insert_writing_tarea1_db(conn, data: dict, module_type_id: int, filename: str) -> int:
    """Insert writing tarea1 data into database"""
    cursor = conn.cursor()
    
    try:
        # Insert into writing_tarea1_set
        cursor.execute("""
            INSERT INTO writing_tarea1_set 
            (title, situation, task_instructions, word_limit, text_type, register, reminders, audio_url, module_type_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING tarea1_set_id;
        """, (
            data.get('title') or filename,
            data.get('situation'),
            data.get('task_instructions'),
            data.get('word_limit'),
            data.get('text_type'),
            data.get('register'),
            data.get('reminders'),
            data.get('audio_url'),
            module_type_id
        ))
        
        tarea1_set_id = cursor.fetchone()[0]
        
        # Insert solution if exists
        if data.get('solution_text'):
            cursor.execute("""
                INSERT INTO writing_tarea1_solution 
                (tarea1_set_id, solution_text)
                VALUES (%s, %s);
            """, (tarea1_set_id, data.get('solution_text')))
        
        conn.commit()
        return tarea1_set_id
        
    except Exception as e:
        conn.rollback()
        raise
    finally:
        cursor.close()


@app.post("/recreate-and-parse-writing-tarea1")
async def recreate_and_parse_writing_tarea1():
    """
    Drop existing writing_tarea1 tables, recreate them, and parse all documents.
    This is the complete deployment endpoint.
    """
    try:
        from parse_writing_tarea1 import recreate_tables, parse_all_documents
        
        folder_path = "Writing_Tarea_1"
        module_type_id = 1
        
        if not os.path.exists(folder_path):
            return {
                "success": False,
                "error": f"Folder '{folder_path}' not found"
            }
        
        # Step 1: Recreate tables
        conn = get_db_connection()
        recreate_tables(conn)
        conn.close()
        
        # Step 2: Parse and insert all documents
        results = parse_all_documents(
            folder_path=folder_path,
            module_type_id=module_type_id,
            limit=None,  # Process all files
            dry_run=False  # Insert into database
        )
        
        return {
            "success": True,
            "message": "Tables recreated and all documents parsed",
            "results": {
                "total": results['total'],
                "successful": results['successful'],
                "failed": results['failed'],
                "errors": results['errors'][:10] if results['errors'] else []
            }
        }
        
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


@app.post("/add-bucket-url-column-writing-tarea1")
async def add_bucket_url_column_writing_tarea1():
    """
    Add bucket_url column to writing_tarea1_set table if it doesn't exist.
    Safe to run multiple times.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Check if column exists
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='writing_tarea1_set' 
            AND column_name='bucket_url';
        """)
        
        column_exists = cursor.fetchone() is not None
        
        if column_exists:
            cursor.close()
            conn.close()
            return {
                "success": True,
                "message": "bucket_url column already exists",
                "action": "none"
            }
        
        # Add the column
        cursor.execute("""
            ALTER TABLE writing_tarea1_set 
            ADD COLUMN bucket_url TEXT;
        """)
        
        cursor.execute("""
            COMMENT ON COLUMN writing_tarea1_set.bucket_url 
            IS 'Google Cloud Storage bucket URL after migration (gs://...)';
        """)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": "bucket_url column added successfully",
            "action": "added"
        }
        
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


@app.get("/writing-tarea1-status")
async def get_writing_tarea1_status():
    """Get status of Writing Tarea 1 parsing and audio migration"""
    try:
        folder_path = "Writing_Tarea_1"
        
        # Count files in folder
        if os.path.exists(folder_path):
            total_files = len([f for f in os.listdir(folder_path) if f.endswith('.docx')])
        else:
            total_files = 0
        
        # Count records in database
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Check if bucket_url column exists
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='writing_tarea1_set' 
            AND column_name='bucket_url';
        """)
        
        bucket_url_exists = cursor.fetchone() is not None
        
        if bucket_url_exists:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_records,
                    COUNT(situation) as with_situation,
                    COUNT(task_instructions) as with_task,
                    COUNT(audio_url) as with_audio_url,
                    COUNT(bucket_url) as with_bucket_url,
                    COUNT(CASE WHEN audio_url IS NOT NULL AND bucket_url IS NULL THEN 1 END) as pending_migration
                FROM writing_tarea1_set;
            """)
        else:
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_records,
                    COUNT(situation) as with_situation,
                    COUNT(task_instructions) as with_task,
                    COUNT(audio_url) as with_audio_url,
                    0 as with_bucket_url,
                    COUNT(audio_url) as pending_migration
                FROM writing_tarea1_set;
            """)
        
        db_stats = dict(cursor.fetchone())
        db_stats['bucket_url_column_exists'] = bucket_url_exists
        
        cursor.execute("""
            SELECT COUNT(*) as with_solution
            FROM writing_tarea1_solution;
        """)
        
        db_stats['with_solution'] = cursor.fetchone()['with_solution']
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "folder": {
                "total_files": total_files,
                "folder_exists": os.path.exists(folder_path)
            },
            "database": db_stats,
            "pending": total_files - db_stats['total_records']
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/migrate-audio-files-writing-tarea1")
async def migrate_audio_files_writing_tarea1(batch_size: int = 10):
    """
    Migrate audio files from Google Drive to GCS bucket for Writing Tarea 1.
    Similar to listening tarea migrations.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get records with audio_url but no bucket_url (or Drive URLs in bucket_url)
        cursor.execute("""
            SELECT tarea1_set_id, audio_url, bucket_url
            FROM writing_tarea1_set
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL 
                 OR bucket_url = '' 
                 OR bucket_url LIKE '%drive.google%' 
                 OR bucket_url NOT LIKE 'gs://%')
            LIMIT %s;
        """, (batch_size,))
        
        records = cursor.fetchall()
        
        if not records:
            cursor.close()
            conn.close()
            return {
                "success": True,
                "message": "No files to migrate",
                "migrated": 0,
                "failed": 0
            }
        
        migrated = 0
        failed = 0
        errors = []
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        
        for record in records:
            tarea1_set_id = record['tarea1_set_id']
            audio_url = record['audio_url']
            
            try:
                # Download from Google Drive
                file_id = extract_google_drive_id(audio_url)
                if not file_id:
                    raise Exception(f"Could not extract file ID from URL: {audio_url}")
                
                download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
                response = requests.get(download_url, timeout=30)
                
                if response.status_code != 200:
                    raise Exception(f"Failed to download: HTTP {response.status_code}")
                
                audio_data = response.content
                
                # Upload to GCS bucket
                blob_name = f"writing_tarea1/{tarea1_set_id}.mp3"
                blob = bucket.blob(blob_name)
                blob.upload_from_string(audio_data, content_type='audio/mpeg')
                
                gcs_url = f"gs://{GCS_BUCKET_NAME}/{blob_name}"
                
                # Update database
                cursor.execute("""
                    UPDATE writing_tarea1_set
                    SET bucket_url = %s
                    WHERE tarea1_set_id = %s;
                """, (gcs_url, tarea1_set_id))
                
                conn.commit()
                migrated += 1
                
            except Exception as e:
                failed += 1
                errors.append({
                    "id": tarea1_set_id,
                    "url": audio_url,
                    "error": str(e)
                })
                conn.rollback()
        
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "migrated": migrated,
            "failed": failed,
            "errors": errors[:5]  # Return first 5 errors
        }
        
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


@app.get("/start-migration-writing-tarea1")
async def start_auto_migration_writing_tarea1():
    """
    Auto-migrate all audio files for Writing Tarea 1.
    Processes in batches until all files are migrated.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Check if bucket_url column exists
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='writing_tarea1_set' 
            AND column_name='bucket_url';
        """)
        
        column_exists = cursor.fetchone() is not None
        
        if not column_exists:
            cursor.close()
            conn.close()
            return {
                "success": False,
                "error": "bucket_url column does not exist. Please run POST /add-bucket-url-column-writing-tarea1 first or recreate the tables."
            }
        
        # Count total files to migrate
        cursor.execute("""
            SELECT COUNT(*) as total
            FROM writing_tarea1_set
            WHERE audio_url IS NOT NULL 
            AND (bucket_url IS NULL 
                 OR bucket_url = '' 
                 OR bucket_url LIKE '%drive.google%' 
                 OR bucket_url NOT LIKE 'gs://%');
        """)
        
        result = cursor.fetchone()
        if not result:
            cursor.close()
            conn.close()
            return {
                "success": False,
                "error": "No result from count query - table may not exist"
            }
        
        total_to_migrate = result['total']
        cursor.close()
        conn.close()
        
        if total_to_migrate == 0:
            return {
                "success": True,
                "message": "All audio files already migrated",
                "total": 0,
                "migrated": 0,
                "failed": 0
            }
        
        # Process in batches
        total_migrated = 0
        total_failed = 0
        all_errors = []
        batch_size = 10
        
        while True:
            result = await migrate_audio_files_writing_tarea1(batch_size=batch_size)
            
            if not result.get('success'):
                return result
            
            total_migrated += result.get('migrated', 0)
            total_failed += result.get('failed', 0)
            all_errors.extend(result.get('errors', []))
            
            # Break if no more files to process
            if result.get('migrated', 0) == 0 and result.get('failed', 0) == 0:
                break
        
        return {
            "success": True,
            "message": f"Migration completed for Writing Tarea 1",
            "total": total_to_migrate,
            "migrated": total_migrated,
            "failed": total_failed,
            "errors": all_errors[:10]  # Return first 10 errors
        }
        
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }


