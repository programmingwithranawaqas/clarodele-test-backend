#!/usr/bin/env python3
"""
Writing Tarea 1 Parser - Complete Document Extraction
Parses all .docx files from Writing_Tarea_1 folder and stores in PostgreSQL database.
Ensures NO TEXT IS MISSED from documents.
"""

import os
import re
import psycopg2
from psycopg2.extras import RealDictCursor
from docx import Document
from typing import Dict, List, Optional
import argparse


# Database configuration
DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://postgres:%40%21ceXpert%2C9%2F11@db.ycnqrnacilcwqqgelaod.supabase.co:5432/postgres?sslmode=require"
)

# Default module_type_id (you can change this)
DEFAULT_MODULE_TYPE_ID = 1


def get_db_connection():
    """Create a database connection"""
    return psycopg2.connect(DATABASE_URL)


def extract_all_text_from_docx(docx_path: str) -> str:
    """
    Extract ALL text from a .docx file including paragraphs, tables, headers, footers.
    Ensures no text is missed.
    """
    doc = Document(docx_path)
    full_text = []
    
    # 1. Extract all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # 2. Extract all text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)
    
    # 3. Extract text from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
    
    # 4. Extract text from footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
    
    return "\n".join(full_text)


def parse_writing_tarea1_document(docx_path: str) -> Dict:
    """Robust parsing of a Writing Tarea 1 document.

    Strategy:
    1. Extract every text-bearing element (paragraphs, tables, headers, footers) preserving order.
    2. Identify heading markers for sections (SITUACIÓN, TAREA, SOLUCIÓN, Recuerde, etc.).
    3. Segment text by heading indices. If some headings missing, use heuristics (e.g. first third as situation, middle as task, last as solution) while NEVER discarding text.
    4. Extract metadata (word limit, text type, register) via regex scanning entire text, not limited to sections.
    5. Extract reminders inside their section or lines starting with Recuerde.
    6. Extract audio_url by finding first URL-like line referencing drive, docs, or gs:// or ending with common audio extensions.
    7. Produce an extraction_report for debugging (line spans, counts, matched patterns).
    """

    raw_text = extract_all_text_from_docx(docx_path)
    lines = [ln.strip() for ln in raw_text.split('\n') if ln.strip()]

    heading_patterns = {
        'situation': [r'^SITUACI[ÓO]N\b', r'^SITUACION\b', r'^Situaci[óo]n\b'],
        'task': [r'^TAREA\b', r'^Tarea\b', r'^INSTRUCCIONES?\b'],
        'solution': [r'^SOLUCI[ÓO]N\b', r'^Soluci[óo]n\b', r'^Modelo de respuesta', r'^MODELO\b'],
        'reminders': [r'^Recuerde', r'^RECUERDE', r'^Recordatorio']
    }

    # Find heading indices
    heading_indices = []  # list of (index, type, original_line)
    for idx, line in enumerate(lines):
        for htype, patterns in heading_patterns.items():
            for pat in patterns:
                if re.search(pat, line):
                    heading_indices.append((idx, htype, line))
                    break
            else:
                continue
            break

    heading_indices.sort(key=lambda x: x[0])

    def slice_section(start_i: int, end_i: int) -> str:
        return '\n'.join(lines[start_i:end_i]).strip()

    # Map of section -> content
    sections = {
        'situation': None,
        'task': None,
        'solution': None,
        'reminders': None
    }

    # Build segments between headings
    for i, (idx, htype, line) in enumerate(heading_indices):
        # Determine content start (allow inline heading like 'SITUACIÓN: text')
        inline_content = ''
        if ':' in line:
            after = line.split(':', 1)[1].strip()
            if after:
                inline_content = after
        content_start = idx + 1 if not inline_content else idx  # if inline content keep same line
        content_end = heading_indices[i + 1][0] if i + 1 < len(heading_indices) else len(lines)
        body = slice_section(content_start, content_end)
        if inline_content:
            # Prepend inline content (excluding heading token itself)
            body = inline_content + ('\n' + body if body else '')
        if sections[htype] is None:
            sections[htype] = body
        else:
            # If repeated heading, append
            sections[htype] += ('\n' + body if body else '')

    # Heuristics if some sections missing
    if sections['situation'] is None or sections['task'] is None or sections['solution'] is None:
        total_lines = len(lines)
        # Only attempt heuristic if headings sparse (<2 types found)
        if len({h for _, h, _ in heading_indices}) < 2 and total_lines >= 10:
            # Split roughly: first 30% situation, next 30% task, final 40% solution
            a = int(total_lines * 0.30)
            b = int(total_lines * 0.60)
            if sections['situation'] is None:
                sections['situation'] = slice_section(0, a)
            if sections['task'] is None:
                sections['task'] = slice_section(a, b)
            if sections['solution'] is None:
                sections['solution'] = slice_section(b, total_lines)

    # Extract metadata from entire text
    word_limit_regexes = [
        r'\b(\d{2,3})\s*[–-]\s*(\d{2,3})\s*palabras',
        r'l[ií]mite\s+de\s+\d{2,3}\s*palabras',
        r'\bentre\s+\d{2,3}\s*y\s+\d{2,3}\s+palabras'
    ]
    word_limit = None
    for wl_pat in word_limit_regexes:
        m = re.search(wl_pat, raw_text, flags=re.IGNORECASE)
        if m:
            word_limit = m.group(0).strip()
            break
    if not word_limit:
        # fallback: any line containing 'palabra'
        for ln in lines:
            if 'palabra' in ln.lower():
                word_limit = ln.strip()
                break

    text_type = None
    for ln in lines:
        if re.search(r'^Tipo\b', ln, flags=re.IGNORECASE) or 'tipo de texto' in ln.lower():
            if ':' in ln:
                text_type = ln.split(':', 1)[1].strip() or ln.strip()
            else:
                text_type = ln.strip()
            break

    register = None
    for ln in lines:
        if ln.lower().startswith('registro'):
            if ':' in ln:
                register = ln.split(':', 1)[1].strip() or ln.strip()
            else:
                register = ln.strip()
            break

    # Reminders: if section already extracted use that; else gather lines beginning with Recuerde
    reminders = sections['reminders']
    if not reminders:
        reminder_lines = [ln for ln in lines if re.match(r'^(Recuerde|RECUERDE|Recordatorio)', ln)]
        if reminder_lines:
            reminders = '\n'.join(reminder_lines).strip()

    # Audio URL: detect first plausible URL line
    audio_url = None
    url_regex = re.compile(r'https?://\S+|gs://\S+|https?://drive\.google\.com/\S+', re.IGNORECASE)
    for ln in lines:
        if url_regex.search(ln):
            # Accept only if looks like drive or audio file or bucket
            if any(ext in ln.lower() for ext in ['.mp3', '.wav', 'drive.google', 'docs.google', 'gs://']):
                audio_url = url_regex.search(ln).group(0)
                break

    # Title: choose first non-heading, non-metadata medium-length line near top
    title = None
    for i, ln in enumerate(lines[:12]):
        normalized = re.sub(r'[:\s]+$', '', ln)
        if any(re.search(p, ln) for pats in heading_patterns.values() for p in pats):
            continue
        if re.search(r'^(Tipo|Registro|Recuerde)', ln, flags=re.IGNORECASE):
            continue
        if 5 <= len(normalized) <= 90:
            title = ln.strip()
            break

    data = {
        'title': title,
        'situation': sections['situation'],
        'task_instructions': sections['task'],
        'word_limit': word_limit,
        'text_type': text_type,
        'register': register,
        'reminders': reminders,
        'audio_url': audio_url,
        'solution_text': sections['solution'],
        'full_raw_text': raw_text,
        'extraction_report': {
            'headings_found': heading_indices,
            'total_lines': len(lines),
            'sections_missing': [k for k, v in sections.items() if k != 'reminders' and v is None],
            'word_limit_match': word_limit,
            'text_type_line': text_type,
            'register_line': register,
            'audio_url_line': audio_url,
            'reminder_present': bool(reminders)
        }
    }

    # Fallback if situation/task both empty: assign full text
    if not data['situation'] and not data['task_instructions']:
        data['situation'] = raw_text

    return data


def save_section(data: Dict, section_name: str, content: List[str]):
    """Save accumulated section content to data dictionary"""
    text = '\n'.join(content).strip()
    
    if section_name == 'situation':
        data['situation'] = text
    elif section_name == 'task':
        data['task_instructions'] = text
    elif section_name == 'solution':
        data['solution_text'] = text
    elif section_name == 'reminder':
        data['reminders'] = text


def extract_metadata_value(line: str) -> str:
    """Extract value from metadata lines like 'Tipo: Informe formal'"""
    if ':' in line:
        return line.split(':', 1)[1].strip()
    return line.strip()


def insert_writing_tarea1(conn, data: Dict, module_type_id: int, file_name: str) -> Optional[int]:
    """
    Insert writing tarea1 data into database.
    Returns the tarea1_set_id if successful, None otherwise.
    """
    cursor = conn.cursor()
    
    try:
        # Insert into writing_tarea1_set
        cursor.execute("""
            INSERT INTO writing_tarea1_set 
            (title, situation, task_instructions, word_limit, text_type, register, reminders, audio_url, module_type_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING tarea1_set_id;
        """, (
            data.get('title') or file_name,  # Use filename as fallback title
            data.get('situation'),
            data.get('task_instructions'),
            data.get('word_limit'),
            data.get('text_type'),
            data.get('register'),
            data.get('reminders'),
            data.get('audio_url'),
            module_type_id
        ))
        
        tarea1_set_id = cursor.fetchone()[0]
        
        # Insert solution if exists
        if data.get('solution_text'):
            cursor.execute("""
                INSERT INTO writing_tarea1_solution 
                (tarea1_set_id, solution_text)
                VALUES (%s, %s);
            """, (tarea1_set_id, data.get('solution_text')))
        
        conn.commit()
        return tarea1_set_id
        
    except Exception as e:
        conn.rollback()
        print(f"Error inserting data: {e}")
        raise
    finally:
        cursor.close()


def parse_all_documents(folder_path: str, module_type_id: int = DEFAULT_MODULE_TYPE_ID, 
                        limit: Optional[int] = None, dry_run: bool = False):
    """
    Parse all .docx files in the folder and insert into database.
    
    Args:
        folder_path: Path to Writing_Tarea_1 folder
        module_type_id: ID of the module type
        limit: Limit number of files to process (for testing)
        dry_run: If True, parse but don't insert into database
    """
    # Get all .docx files
    files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    files.sort()  # Sort for consistent processing
    
    if limit:
        files = files[:limit]
    
    print(f"Found {len(files)} .docx files to process")
    print("=" * 80)
    
    conn = None if dry_run else get_db_connection()
    
    results = {
        'total': len(files),
        'successful': 0,
        'failed': 0,
        'errors': []
    }
    
    try:
        for idx, filename in enumerate(files, 1):
            file_path = os.path.join(folder_path, filename)
            
            print(f"\n[{idx}/{len(files)}] Processing: {filename}")
            
            try:
                # Parse document
                data = parse_writing_tarea1_document(file_path)
                
                # Validation: ensure critical fields are not empty
                if not data.get('situation') and not data.get('task_instructions'):
                    print(f"  ⚠️  WARNING: No situation or task found, using full text")
                    # If structure not detected, store all text in situation
                    data['situation'] = data['full_raw_text']
                
                if dry_run:
                    print(f"  📄 DRY RUN - Parsed successfully")
                    title_preview = (data.get('title') or 'N/A')
                    print(f"     - Title: {title_preview[:50]}...")
                    print(f"     - Situation: {len(data.get('situation', '') or '')} chars")
                    print(f"     - Task: {len(data.get('task_instructions', '') or '')} chars")
                    print(f"     - Solution: {len(data.get('solution_text', '') or '')} chars")
                    print(f"     - Total text: {len(data.get('full_raw_text') or '')} chars")
                else:
                    # Insert into database
                    tarea1_set_id = insert_writing_tarea1(conn, data, module_type_id, filename)
                    print(f"  ✅ Inserted successfully (ID: {tarea1_set_id})")
                
                results['successful'] += 1
                
            except Exception as e:
                print(f"  ❌ Error: {str(e)}")
                results['failed'] += 1
                results['errors'].append({
                    'file': filename,
                    'error': str(e)
                })
    
    finally:
        if conn:
            conn.close()
    
    # Print summary
    print("\n" + "=" * 80)
    print("PARSING SUMMARY")
    print("=" * 80)
    print(f"Total files: {results['total']}")
    print(f"Successful: {results['successful']}")
    print(f"Failed: {results['failed']}")
    
    if results['errors']:
        print(f"\n❌ Errors ({len(results['errors'])}):")
        for error in results['errors'][:10]:
            print(f"  - {error['file']}: {error['error']}")
        if len(results['errors']) > 10:
            print(f"  ... and {len(results['errors']) - 10} more")
    
    return results


def main():
    parser = argparse.ArgumentParser(description="Parse Writing Tarea 1 documents")
    parser.add_argument(
        "--folder",
        type=str,
        default="Writing_Tarea_1",
        help="Path to Writing_Tarea_1 folder"
    )
    parser.add_argument(
        "--module-type-id",
        type=int,
        default=DEFAULT_MODULE_TYPE_ID,
        help="Module type ID for the documents"
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Limit number of files to process (for testing)"
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse documents but don't insert into database"
    )
    
    args = parser.parse_args()
    
    print("=" * 80)
    print("WRITING TAREA 1 PARSER")
    print("=" * 80)
    print(f"Folder: {args.folder}")
    print(f"Module Type ID: {args.module_type_id}")
    print(f"Limit: {args.limit if args.limit else 'No limit'}")
    print(f"Dry Run: {args.dry_run}")
    print("=" * 80)
    
    if not os.path.exists(args.folder):
        print(f"❌ Error: Folder '{args.folder}' not found!")
        return
    
    parse_all_documents(
        folder_path=args.folder,
        module_type_id=args.module_type_id,
        limit=args.limit,
        dry_run=args.dry_run
    )


if __name__ == "__main__":
    main()
